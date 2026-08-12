"""CV parsing, intelligent extraction, embeddings and job↔CV matching.

Design rules from the spec:
  * numerical similarity ALWAYS comes from embeddings, never from the LLM;
  * LLM extraction is optional — a deterministic fallback keeps everything working
    when no API key is set or the model chain is exhausted;
  * embedding backend degrades to a local lexical vectorizer if
    sentence-transformers / model weights are unavailable offline.
"""
from __future__ import annotations

import io
import logging
import math
import re
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

import numpy as np

from app.config import settings
from app.schemas import CVProfile, Job
from app.services.llm_client import LLMUnavailable, llm_client
from app.services.skills_taxonomy import (
    SOFT_SKILLS,
    canonicalize,
    extract_skills_from_text,
)

logger = logging.getLogger(__name__)


# ==========================================================================
# 1. Text extraction from files
# ==========================================================================
class CVParseError(ValueError):
    pass


def extract_text_from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception as exc:  # noqa: BLE001
                logger.warning("PDF page extraction failed: %s", exc)
        return "\n".join(pages).strip()
    except Exception as exc:  # noqa: BLE001
        raise CVParseError(f"Could not read PDF: {exc}") from exc


def extract_text_from_docx(data: bytes) -> str:
    import docx  # python-docx

    try:
        document = docx.Document(io.BytesIO(data))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(p for p in parts if p and p.strip()).strip()
    except Exception as exc:  # noqa: BLE001
        raise CVParseError(f"Could not read DOCX: {exc}") from exc


def extract_cv_text(data: bytes, filename: str) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        text = extract_text_from_pdf(data)
    elif name.endswith((".docx", ".doc")):
        text = extract_text_from_docx(data)
    elif name.endswith((".txt", ".md")):
        text = data.decode("utf-8", errors="ignore")
    else:
        raise CVParseError(
            f"Unsupported file type '{filename}'. Upload a PDF, DOCX or TXT file."
        )
    if len(text.strip()) < 30:
        raise CVParseError(
            "Extracted almost no text. The file may be a scanned image — "
            "please upload a text-based PDF/DOCX."
        )
    return text


# ==========================================================================
# 2. Intelligent CV extraction (LLM first, heuristic fallback)
# ==========================================================================
_EXTRACTION_PROMPT = """Extract structured data from this CV/resume.

Return JSON with EXACTLY these keys:
{{
  "skills": [string],            // concrete skills, max 30
  "job_titles": [string],        // roles held or targeted, max 8
  "experience": [string],        // "Title at Company (dates)" lines, max 10
  "education": [string],         // degrees/institutions, max 6
  "technologies": [string],      // tools/frameworks/languages, max 30
  "certifications": [string],    // max 8
  "years_experience": number,    // total professional years, 0 if unclear
  "summary": string              // 2 sentence professional summary
}}

Use the CV's own wording. Never invent anything not present. Empty list if unknown.

CV:
---
{cv_text}
---"""

_TITLE_HINTS = [
    "engineer", "developer", "analyst", "manager", "designer", "scientist",
    "specialist", "consultant", "architect", "administrator", "lead",
    "intern", "director", "officer", "accountant", "marketer", "technician",
]
_EDU_HINTS = [
    "bachelor", "master", "phd", "b.sc", "m.sc", "bsc", "msc", "university",
    "faculty", "degree", "diploma", "بكالوريوس", "ماجستير", "جامعة", "كلية",
]
_CERT_HINTS = [
    "certified", "certificate", "certification", "aws certified", "pmp",
    "scrum master", "nanodegree", "coursera", "udacity", "شهادة",
]


def _dedupe_keep_order(items: Iterable[str], limit: int) -> List[str]:
    out: List[str] = []
    seen = set()
    for item in items:
        cleaned = re.sub(r"\s+", " ", str(item)).strip(" -•\t")
        if not cleaned or len(cleaned) > 120:
            continue
        key = cleaned.lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(cleaned)
        if len(out) >= limit:
            break
    return out


def heuristic_profile(cv_text: str) -> CVProfile:
    """Deterministic extraction — no LLM required."""
    lines = [l.strip() for l in cv_text.splitlines() if l.strip()]
    lower = cv_text.lower()

    skills = extract_skills_from_text(cv_text, limit=40)

    titles = [
        l for l in lines
        if len(l) < 90 and any(h in l.lower() for h in _TITLE_HINTS)
    ]
    education = [l for l in lines if any(h in l.lower() for h in _EDU_HINTS)]
    certifications = [l for l in lines if any(h in l.lower() for h in _CERT_HINTS)]

    # Experience-ish lines: contain a year range or "at Company"
    experience = [
        l for l in lines
        if re.search(r"(19|20)\d{2}\s*[-–—to]+\s*((19|20)\d{2}|present|current)", l.lower())
    ]

    years = 0.0
    m = re.search(r"(\d{1,2})\+?\s*(?:years|yrs|سنة|سنوات)", lower)
    if m:
        years = float(m.group(1))
    else:
        found_years = [int(y) for y in re.findall(r"(19[89]\d|20[0-4]\d)", cv_text)]
        if len(found_years) >= 2:
            span = max(found_years) - min(found_years)
            years = float(min(span, 45))

    summary = " ".join(lines[:3])[:400] if lines else ""

    return CVProfile(
        skills=_dedupe_keep_order(skills, 30),
        job_titles=_dedupe_keep_order(titles, 8),
        experience=_dedupe_keep_order(experience, 10),
        education=_dedupe_keep_order(education, 6),
        technologies=_dedupe_keep_order(skills, 30),
        certifications=_dedupe_keep_order(certifications, 8),
        years_experience=years,
        summary=summary,
    )


async def extract_cv_profile(cv_text: str) -> Tuple[CVProfile, str]:
    """Returns (profile, method) where method is 'llm' or 'heuristic'."""
    fallback = heuristic_profile(cv_text)
    if not llm_client.enabled:
        return fallback, "heuristic"

    try:
        data = await llm_client.complete_json(
            _EXTRACTION_PROMPT.format(cv_text=cv_text[:12000]),
            system="You are an expert resume parser. Output JSON only.",
            max_tokens=1600,
        )
        if not isinstance(data, dict):
            raise ValueError("LLM returned non-object JSON")

        def _list(key: str, limit: int) -> List[str]:
            value = data.get(key) or []
            if isinstance(value, str):
                value = [value]
            return _dedupe_keep_order([str(v) for v in value if v], limit)

        try:
            years = float(data.get("years_experience") or 0)
        except (TypeError, ValueError):
            years = fallback.years_experience

        profile = CVProfile(
            skills=_list("skills", 30) or fallback.skills,
            job_titles=_list("job_titles", 8) or fallback.job_titles,
            experience=_list("experience", 10) or fallback.experience,
            education=_list("education", 6) or fallback.education,
            technologies=_list("technologies", 30) or fallback.technologies,
            certifications=_list("certifications", 8) or fallback.certifications,
            years_experience=years or fallback.years_experience,
            summary=str(data.get("summary") or fallback.summary)[:600],
        )
        # Union with deterministic scan so we never lose obvious skills.
        merged = _dedupe_keep_order(
            [canonicalize(s) for s in profile.skills] + fallback.skills, 40
        )
        profile.skills = merged
        return profile, "llm"
    except (LLMUnavailable, ValueError, KeyError, TypeError) as exc:
        logger.warning("LLM CV extraction failed (%s) — using heuristic fallback", exc)
        return fallback, "heuristic"


# ==========================================================================
# 3. Embeddings
# ==========================================================================
class LexicalEmbedder:
    """Offline fallback: hashed bag-of-words + char-trigram vectors, L2 normalised.

    Not as semantic as a transformer, but deterministic, dependency-free and
    good enough to rank job descriptions against a CV when model weights
    cannot be downloaded.
    """

    name = "lexical-fallback"
    dim = 2048

    _token_re = re.compile(r"[a-z0-9\u0600-\u06FF#+.]{2,}")

    def _tokens(self, text: str) -> List[str]:
        low = text.lower()
        words = self._token_re.findall(low)
        grams = [f"{a}_{b}" for a, b in zip(words, words[1:])]
        return words + grams

    def encode_one(self, text: str) -> np.ndarray:
        vec = np.zeros(self.dim, dtype=np.float32)
        tokens = self._tokens(text or "")
        if not tokens:
            return vec
        counts: Dict[int, float] = {}
        for token in tokens:
            idx = hash(token) % self.dim
            counts[idx] = counts.get(idx, 0.0) + 1.0
        for idx, count in counts.items():
            vec[idx] = 1.0 + math.log(count)  # sublinear tf
        norm = np.linalg.norm(vec)
        return vec / norm if norm else vec

    def encode(self, texts: Sequence[str]) -> np.ndarray:
        return np.vstack([self.encode_one(t) for t in texts])


class EmbeddingService:
    """sentence-transformers when available, lexical fallback otherwise."""

    def __init__(self) -> None:
        self._model = None
        self._backend = "uninitialised"

    def _load(self) -> None:
        if self._model is not None:
            return
        try:
            from sentence_transformers import SentenceTransformer

            self._model = SentenceTransformer(settings.embedding_model)
            self._backend = f"sentence-transformers:{settings.embedding_model}"
            logger.info("Loaded embedding model %s", settings.embedding_model)
            return
        except Exception as exc:  # noqa: BLE001 - offline / no weights / no package
            if not settings.embedding_fallback_enabled:
                raise
            logger.warning(
                "sentence-transformers unavailable (%s) — using lexical fallback", exc
            )
            self._model = LexicalEmbedder()
            self._backend = LexicalEmbedder.name

    @property
    def backend(self) -> str:
        if self._model is None:
            self._load()
        return self._backend

    def encode(self, texts: Sequence[str]) -> np.ndarray:
        self._load()
        if not texts:
            return np.zeros((0, 1), dtype=np.float32)
        vectors = self._model.encode(list(texts))
        vectors = np.asarray(vectors, dtype=np.float32)
        norms = np.linalg.norm(vectors, axis=1, keepdims=True)
        norms[norms == 0] = 1.0
        return vectors / norms


embedding_service = EmbeddingService()


def cosine_similarity(a: np.ndarray, b: np.ndarray) -> float:
    denom = float(np.linalg.norm(a) * np.linalg.norm(b))
    if denom == 0:
        return 0.0
    return float(np.dot(a, b) / denom)


def similarity_to_percentage(similarity: float, backend: str) -> float:
    """Calibrate raw cosine into a human-readable 0–100 match score.

    Transformer cosines for related job text cluster in ~0.2–0.8; lexical
    overlap cosines are lower, so each backend gets its own linear stretch.
    """
    similarity = max(0.0, min(1.0, similarity))
    if backend.startswith("sentence-transformers"):
        lo, hi = 0.05, 0.80
    else:
        lo, hi = 0.02, 0.55
    scaled = (similarity - lo) / (hi - lo)
    return round(max(0.0, min(1.0, scaled)) * 100, 1)


# ==========================================================================
# 4. Matching
# ==========================================================================
def _cv_document(profile: CVProfile, raw_text: str) -> str:
    parts = [
        " ".join(profile.job_titles),
        " ".join(profile.skills),
        " ".join(profile.technologies),
        profile.summary,
        " ".join(profile.experience),
        raw_text[:4000],
    ]
    return "\n".join(p for p in parts if p)


def _job_document(job: Job) -> str:
    return "\n".join(
        p for p in [job.title, job.company, job.location, job.job_type or "", job.description] if p
    )


def compare_skills(
    cv_skills: Sequence[str], job_text: str
) -> Tuple[List[str], List[str]]:
    """Return (matching, missing) skills for a job, using the taxonomy scan."""
    job_skills = extract_skills_from_text(job_text, limit=40)
    cv_canonical = {canonicalize(s).lower() for s in cv_skills if s}
    # Also allow substring matches for skills not in the taxonomy.
    cv_raw = {s.lower().strip() for s in cv_skills if s}

    matching, missing = [], []
    for skill in job_skills:
        key = skill.lower()
        if key in cv_canonical or any(key in raw or raw in key for raw in cv_raw if len(raw) > 2):
            matching.append(skill)
        else:
            missing.append(skill)
    # Soft skills last in the missing list — they're rarely the real gap.
    missing.sort(key=lambda s: (s in SOFT_SKILLS, s))
    return matching, missing


def build_explanation(
    score: float, matching: Sequence[str], missing: Sequence[str], job: Job
) -> str:
    bits: List[str] = []
    if score >= 75:
        bits.append("Strong match")
    elif score >= 50:
        bits.append("Good match")
    elif score >= 30:
        bits.append("Partial match")
    else:
        bits.append("Weak match")
    if matching:
        bits.append(f"shares {len(matching)} skill(s): {', '.join(list(matching)[:6])}")
    else:
        bits.append("no directly overlapping skills detected")
    if missing:
        hard = [m for m in missing if m not in SOFT_SKILLS]
        if hard:
            bits.append(f"gaps: {', '.join(hard[:5])}")
    if job.location:
        bits.append(f"location: {job.location}")
    return " — ".join(bits) + "."


def rank_jobs(
    jobs: List[Job], profile: CVProfile, cv_raw_text: str = ""
) -> List[Job]:
    """Score every job by CV↔job embedding cosine similarity, then sort desc."""
    if not jobs:
        return []

    cv_doc = _cv_document(profile, cv_raw_text)
    job_docs = [_job_document(j) for j in jobs]

    try:
        vectors = embedding_service.encode([cv_doc] + job_docs)
        cv_vec, job_vecs = vectors[0], vectors[1:]
        backend = embedding_service.backend
    except Exception as exc:  # noqa: BLE001
        logger.error("Embedding failed (%s) — falling back to skill overlap only", exc)
        cv_vec, job_vecs, backend = None, None, "unavailable"

    for idx, job in enumerate(jobs):
        matching, missing = compare_skills(profile.skills, _job_document(job))
        if cv_vec is not None:
            similarity = cosine_similarity(cv_vec, job_vecs[idx])
            score = similarity_to_percentage(similarity, backend)
        else:
            total = len(matching) + len(missing)
            score = round((len(matching) / total) * 100, 1) if total else 0.0

        # Small deterministic bonus for explicit skill overlap (never LLM-invented).
        if matching:
            score = min(100.0, score + min(len(matching) * 1.5, 12.0))

        job.match_score = round(score, 1)
        job.matching_skills = matching[:15]
        job.missing_skills = missing[:15]
        job.match_explanation = build_explanation(job.match_score, matching, missing, job)

    return sorted(jobs, key=lambda j: (j.match_score or 0.0), reverse=True)
